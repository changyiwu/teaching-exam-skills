#!/usr/bin/env python
"""
generate_exam_docx.py — 段考試卷 / 情境非選題 Word 產生器

同時支援兩種 JSON schema，依欄位自動判斷：

  A. 段考卷（jh-math-exam）          有 mc_questions 或 open_questions
     產出：題目卷、答案卷（教師版）、雙向細目表〔、審題報告〕
  B. 情境非選題（jh-math-context-questions）  有 questions
     產出：單一份 Word（題目 + 解析）

版面規格見 references/exam-format.md 與 references/shuangxiang-table.md。

用法：
    python generate_exam_docx.py <exam_data.json> <輸出目錄或 .docx>
                                 [--blueprint-only] [--figures-dir DIR]

  --blueprint-only  只產雙向細目表（＋審題報告），用於審題模式
  --figures-dir     幾何圖 PNG 所在目錄；有此參數時會依題目的 geometry
                    欄位自動把圖插在正確位置（不需事後用正規表達式回頭找段落）
"""
import argparse
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

sys.path.insert(0, str(Path(__file__).resolve().parent))
from math_omml import add_rich_text, plain_text, set_run_font  # noqa: E402

BODY_FONT = "標楷體"
NOTE_FONT = "新細明體"
BODY_SIZE = 12

# 內部四級 → 學校空白表的欄名（見 references/shuangxiang-table.md）
BLOOM_COLUMNS = ["記憶", "了解", "應用", "思考判斷\n（分析、綜合、評鑑）"]
BLOOM_INTERNAL = ["第1級（記憶）", "第2級（理解）", "第3級（應用）", "第4級（分析）"]
BLOOM_SUGGEST = [10, 20, 40, 30]


# ══════════════════════════════════════════════════════
# 共用工具
# ══════════════════════════════════════════════════════

def norm_bloom(value):
    """各種寫法的 bloom_level → 0..3；判不出來回傳 None。"""
    if value is None:
        return None
    s = str(value)
    m = re.search(r"第\s*([1-4])\s*級", s)
    if m:
        return int(m.group(1)) - 1
    for idx, keys in enumerate((("記憶", "記憶性"), ("理解", "了解"),
                                ("應用", "運用"), ("分析", "思考判斷", "綜合", "評鑑"))):
        if any(k in s for k in keys):
            return idx
    return None


def parse_mc_scoring(text, count):
    """'1~7題每題4分、8~16題每題3分' → {題號: 配分}。解析不出來回傳 {}。"""
    if not text:
        return {}
    points = {}
    for a, b, p in re.findall(r"(\d+)\s*[~～\-–至]\s*(\d+)\s*題[^0-9]*?(\d+)\s*分", str(text)):
        for q in range(int(a), int(b) + 1):
            points[q] = int(p)
    if not points:
        m = re.search(r"每題\s*(\d+)\s*分", str(text))
        if m:
            points = {q: int(m.group(1)) for q in range(1, count + 1)}
    return points


def mc_points(data):
    """回傳 {題號: 配分}：優先用各題自帶的 points，其次解析 mc_scoring。"""
    questions = data.get("mc_questions", []) or []
    explicit = {q["number"]: q["points"] for q in questions if q.get("points")}
    if len(explicit) == len(questions) and questions:
        return explicit
    parsed = parse_mc_scoring(data.get("mc_scoring"), len(questions))
    result = {}
    for q in questions:
        n = q["number"]
        val = explicit.get(n) or parsed.get(n)
        if val is None:
            print(f"[WARN] 選擇題第 {n} 題無配分：請在該題加 points 欄位，"
                  f"或讓 mc_scoring 可被解析（如「1~16題每題4分」）")
            val = 0
        result[n] = val
    return result


def setup_page(doc, width_cm, height_cm, margins_cm, lr=None):
    sec = doc.sections[0]
    sec.page_width = Cm(width_cm)
    sec.page_height = Cm(height_cm)
    sec.top_margin = sec.bottom_margin = Cm(margins_cm)
    sec.left_margin = sec.right_margin = Cm(lr if lr is not None else margins_cm)
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(BODY_SIZE)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    return sec


def para(doc, text="", size=BODY_SIZE, bold=False, align=None,
         font=BODY_FONT, space_after=4, rich=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        if rich:
            add_rich_text(p, text, font=font, size=size, bold=bold)
        else:
            run = p.add_run(str(text))
            set_run_font(run, font, size, bold)
    return p


def set_cell_text(cell, text, size=10, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER,
                  font=BODY_FONT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    for k, line in enumerate(str(text).split("\n")):
        if k:
            p.add_run().add_break()
        if line:
            set_run_font(p.add_run(line), font, size, bold)


def shade_cell(cell, hex_color="D9D9D9"):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def insert_figure(doc, png_path, caption="", width_cm=5.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(png_path), width=Cm(width_cm))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(4)
        run = cp.add_run(caption)
        set_run_font(run, BODY_FONT, 9)
        run.italic = True


class Figures:
    """依 geometry 欄位在建構文件時直接插圖，位置必然正確。"""

    def __init__(self, figures_dir):
        self.dir = Path(figures_dir) if figures_dir else None
        self.missing = []

    def add(self, doc, block, fig_id, width_cm=5.5):
        geo = (block or {}).get("geometry")
        if not geo or not self.dir:
            return False
        png = self.dir / f"{fig_id}.png"
        if not png.exists():
            self.missing.append(fig_id)
            return False
        insert_figure(doc, png, geo.get("caption", ""), width_cm)
        return True


# ══════════════════════════════════════════════════════
# A. 段考卷 — 題目卷
# ══════════════════════════════════════════════════════

def build_exam_paper(data, out_dir, figs):
    doc = Document()
    setup_page(doc, 25.7, 36.4, 1.5)

    # 卷頭（單格表格）
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    grade_cn = str(data.get("grade", "")).replace("年級", "")
    head = (
        f"沙鹿國中{data.get('school_year','')}學年度"
        f"第{data.get('semester','')}學期第{data.get('exam_number','')}次段考"
        f"　數學科　考試卷\n"
        f"{grade_cn}年 ___班　座號:___　姓名:__________\n"
        f"範圍：{data.get('scope','')}"
    )
    set_cell_text(cell, head, size=13, bold=True)
    doc.add_paragraph()

    mc = data.get("mc_questions", []) or []
    opens = data.get("open_questions", []) or []
    pts = mc_points(data)

    # ── 選擇題 ──────────────────────────────────────
    if mc:
        note = data.get("mc_scoring") or f"選擇 {len(mc)} 題，共 {sum(pts.values())} 分。"
        para(doc, f"一、選擇題（{note}）", size=12, bold=True, font=NOTE_FONT)
        for q in mc:
            p = para(doc, "", space_after=2)
            set_run_font(p.add_run("（　　）"), BODY_FONT, BODY_SIZE)
            add_rich_text(p, f"{q['number']}. {q.get('question','')}", size=BODY_SIZE)
            figs.add(doc, q, f"mc_{q['number']}_fig")
            opts = q.get("options", {}) or {}
            texts = [f"({k}) {opts.get(k,'')}" for k in ("A", "B", "C", "D")]
            longest = max((len(plain_text(t)) for t in texts), default=0)
            per_line = 4 if longest <= 12 else (2 if longest <= 28 else 1)
            for i in range(0, 4, per_line):
                op = para(doc, "", space_after=2)
                add_rich_text(op, "　".join(texts[i:i + per_line]), size=BODY_SIZE)
        doc.add_paragraph()

    # ── 非選擇題 ────────────────────────────────────
    if opens:
        total = sum(sq.get("points", 0) for q in opens for sq in q.get("sub_questions", []))
        para(doc, f"二、非選擇題（共 {total} 分，須寫出計算過程）",
             size=12, bold=True, font=NOTE_FONT)
        for q in opens:
            n = q["number"]
            head_p = para(doc, "", space_after=2)
            set_run_font(head_p.add_run(f"{n}.（{q.get('total_points','')}分）"),
                         BODY_FONT, BODY_SIZE, True)
            add_rich_text(head_p, q.get("context", ""), size=BODY_SIZE)
            figs.add(doc, q, f"open_{n}_fig", width_cm=6.0)
            for sq in q.get("sub_questions", []) or []:
                label = sq.get("label", "")
                sp = para(doc, "", space_after=2)
                set_run_font(sp.add_run(f"　{label}（{sq.get('points','')}分）"),
                             BODY_FONT, BODY_SIZE, True)
                add_rich_text(sp, sq.get("question", ""), size=BODY_SIZE)
                figs.add(doc, sq, f"open_{n}_{label.strip('()（）')}_fig", width_cm=5.5)
                for _ in range(3):               # 作答空白
                    para(doc, "", space_after=0)

    path = out_dir / f"{data.get('grade','')}數學第{data.get('exam_number','')}次段考_題目卷.docx"
    doc.save(str(path))
    return path


# ══════════════════════════════════════════════════════
# A. 段考卷 — 答案卷（教師版）
# ══════════════════════════════════════════════════════

def build_answer_paper(data, out_dir):
    doc = Document()
    setup_page(doc, 25.7, 36.4, 1.5)
    para(doc, f"{data.get('grade','')} 數學第{data.get('exam_number','')}次段考"
              f"　答案卷（教師版）", size=16, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, font=NOTE_FONT)
    para(doc, f"範圍：{data.get('scope','')}", size=11,
         align=WD_ALIGN_PARAGRAPH.CENTER, font=NOTE_FONT)
    doc.add_paragraph()

    mc = data.get("mc_questions", []) or []
    pts = mc_points(data)

    if mc:
        para(doc, "一、選擇題答案", size=13, bold=True, font=NOTE_FONT)
        # 每 5 題一組的速查表（照範例卷格式）
        rows = [mc[i:i + 5] for i in range(0, len(mc), 5)]
        quick = doc.add_table(rows=len(rows) * 2, cols=6)
        quick.style = "Table Grid"
        for r, group in enumerate(rows):
            set_cell_text(quick.rows[r * 2].cells[0], "題號", 10, True)
            set_cell_text(quick.rows[r * 2 + 1].cells[0], "答案", 10, True)
            shade_cell(quick.rows[r * 2].cells[0])
            shade_cell(quick.rows[r * 2 + 1].cells[0])
            for c, q in enumerate(group):
                set_cell_text(quick.rows[r * 2].cells[c + 1], str(q["number"]), 10)
                set_cell_text(quick.rows[r * 2 + 1].cells[c + 1], str(q.get("answer", "")), 11, True)
        doc.add_paragraph()

        para(doc, "二、選擇題詳解", size=13, bold=True, font=NOTE_FONT)
        for q in mc:
            p = para(doc, "", space_after=2)
            set_run_font(p.add_run(f"{q['number']}. "), BODY_FONT, BODY_SIZE, True)
            add_rich_text(p, q.get("question", ""), size=BODY_SIZE)
            meta = para(doc, "", space_after=2)
            set_run_font(meta.add_run(
                f"　答案：{q.get('answer','')}　｜　配分：{pts.get(q['number'],0)} 分"
                f"　｜　{q.get('section','')}　｜　{q.get('bloom_level','')}"),
                NOTE_FONT, 10, True)
            if q.get("solution"):
                sp = para(doc, "", space_after=6)
                set_run_font(sp.add_run("　解："), BODY_FONT, 11, True)
                add_rich_text(sp, q["solution"], size=11)
        doc.add_paragraph()

    opens = data.get("open_questions", []) or []
    if opens:
        para(doc, "三、非選擇題參考答案與給分要點", size=13, bold=True, font=NOTE_FONT)
        for q in opens:
            hp = para(doc, "", space_after=2)
            set_run_font(hp.add_run(f"{q['number']}.（{q.get('total_points','')}分）"),
                         BODY_FONT, BODY_SIZE, True)
            add_rich_text(hp, q.get("context", ""), size=BODY_SIZE)
            for sq in q.get("sub_questions", []) or []:
                sp = para(doc, "", space_after=2)
                set_run_font(sp.add_run(f"　{sq.get('label','')}（{sq.get('points','')}分）"),
                             BODY_FONT, 11, True)
                add_rich_text(sp, sq.get("question", ""), size=11)
                ap = para(doc, "", space_after=2)
                set_run_font(ap.add_run("　　答："), BODY_FONT, 11, True)
                add_rich_text(ap, sq.get("answer", ""), size=11)
                if sq.get("solution"):
                    lp = para(doc, "", space_after=2)
                    set_run_font(lp.add_run("　　解："), BODY_FONT, 11, True)
                    add_rich_text(lp, sq["solution"], size=11)
                if sq.get("bloom_level"):
                    para(doc, f"　　（{sq['bloom_level']}）", size=9, font=NOTE_FONT, space_after=6)

    path = out_dir / f"{data.get('grade','')}數學第{data.get('exam_number','')}次段考_答案卷（教師版）.docx"
    doc.save(str(path))
    return path


# ══════════════════════════════════════════════════════
# A. 段考卷 — 雙向細目表
# ══════════════════════════════════════════════════════

def collect_blueprint(data):
    """回傳 {section: {型式: [ [配分,題數] × 4 ]}}，型式為 選擇題/計算與證明題。"""
    stats = {}
    order = []

    def bucket(section, kind):
        section = section or "（未標節次）"
        if section not in stats:
            stats[section] = {k: [[0, 0] for _ in range(4)]
                              for k in ("選擇題", "填充題", "計算與證明題")}
            order.append(section)
        return stats[section][kind]

    pts = mc_points(data)
    for q in data.get("mc_questions", []) or []:
        idx = norm_bloom(q.get("bloom_level"))
        if idx is None:
            print(f"[WARN] 選擇題第 {q['number']} 題的 bloom_level 無法判讀：{q.get('bloom_level')!r}")
            continue
        cell = bucket(q.get("section"), "選擇題")[idx]
        cell[0] += pts.get(q["number"], 0)
        cell[1] += 1

    for q in data.get("open_questions", []) or []:
        for sq in q.get("sub_questions", []) or []:
            idx = norm_bloom(sq.get("bloom_level"))
            if idx is None:
                print(f"[WARN] 非選第 {q['number']} 題 {sq.get('label','')} 的 bloom_level "
                      f"無法判讀：{sq.get('bloom_level')!r}")
                continue
            cell = bucket(sq.get("section") or q.get("section"), "計算與證明題")[idx]
            cell[0] += sq.get("points", 0)
            cell[1] += 1

    return stats, order


def build_blueprint(data, out_dir):
    doc = Document()
    setup_page(doc, 21.0, 29.7, 1.5, lr=2.0)

    para(doc, f"台中市沙鹿國中 {data.get('school_year','')} 學年度"
              f"第 {data.get('semester','')} 學期第 {data.get('exam_number','')} 次段考"
              f"試卷雙向細目分析表",
         size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    # 表一：基本資料
    info = doc.add_table(rows=2, cols=4)
    info.style = "Table Grid"
    grade_cn = str(data.get("grade", "")).replace("年級", "")
    for r, (a, b, c, d) in enumerate([
        ("科目名稱", "數學科", "使用年級", f"國中 {grade_cn} 年級"),
        ("命題教師", data.get("teacher", ""), "版本", f"{data.get('version','')} 版"),
    ]):
        for col, val in enumerate((a, b, c, d)):
            set_cell_text(info.rows[r].cells[col], val, 11, bold=(col % 2 == 0))
            if col % 2 == 0:
                shade_cell(info.rows[r].cells[col])
    doc.add_paragraph()

    stats, order = collect_blueprint(data)
    kinds = ("選擇題", "填充題", "計算與證明題", "小計")
    n_rows = 2 + len(order) * 4 + 1
    tb = doc.add_table(rows=n_rows, cols=11)
    tb.style = "Table Grid"

    # 表頭兩列
    set_cell_text(tb.rows[0].cells[0], "教材內容", 10, True)
    set_cell_text(tb.rows[0].cells[1], "試題型式", 10, True)
    for i, name in enumerate(BLOOM_COLUMNS):
        set_cell_text(tb.rows[0].cells[2 + i * 2], name, 9, True)
        tb.rows[0].cells[2 + i * 2].merge(tb.rows[0].cells[3 + i * 2])
        set_cell_text(tb.rows[1].cells[2 + i * 2], "配分", 9, True)
        set_cell_text(tb.rows[1].cells[3 + i * 2], "題數", 9, True)
    set_cell_text(tb.rows[0].cells[10], "合計\n(配分)", 10, True)
    tb.rows[0].cells[0].merge(tb.rows[1].cells[0])
    tb.rows[0].cells[1].merge(tb.rows[1].cells[1])
    tb.rows[0].cells[10].merge(tb.rows[1].cells[10])
    for c in range(11):
        shade_cell(tb.rows[0].cells[c])
        shade_cell(tb.rows[1].cells[c])

    grand = [[0, 0] for _ in range(4)]
    for s_i, section in enumerate(order):
        base = 2 + s_i * 4
        set_cell_text(tb.rows[base].cells[0], section, 10)
        tb.rows[base].cells[0].merge(tb.rows[base + 3].cells[0])
        subtotal = [[0, 0] for _ in range(4)]
        for k_i, kind in enumerate(kinds):
            row = tb.rows[base + k_i]
            set_cell_text(row.cells[1], kind, 9, bold=(kind == "小計"))
            values = subtotal if kind == "小計" else stats[section][kind]
            line_total = 0
            for b in range(4):
                p, c = values[b]
                if kind != "小計":
                    subtotal[b][0] += p
                    subtotal[b][1] += c
                    grand[b][0] += p
                    grand[b][1] += c
                set_cell_text(row.cells[2 + b * 2], p if p else "", 10)
                set_cell_text(row.cells[3 + b * 2], c if c else "", 10)
                line_total += p
            set_cell_text(row.cells[10], line_total if line_total else "", 10,
                          bold=(kind == "小計"))
            if kind == "小計":
                for c in range(11):
                    shade_cell(row.cells[c], "F2F2F2")

    last = tb.rows[n_rows - 1]
    set_cell_text(last.cells[0], "合計", 10, True)
    last.cells[0].merge(last.cells[1])
    for b in range(4):
        set_cell_text(last.cells[2 + b * 2], grand[b][0] or "", 10, True)
        set_cell_text(last.cells[3 + b * 2], grand[b][1] or "", 10, True)
    total = sum(g[0] for g in grand)
    set_cell_text(last.cells[10], total, 11, True)
    for c in range(11):
        shade_cell(last.cells[c])

    doc.add_paragraph()
    para(doc, "【說明】", size=12, bold=True, font=NOTE_FONT, space_after=2)
    for line in ("請命題教師依照每一試題的認知層次，分類填入上表中。",
                 "試題型式請依照試卷編排順序填寫。",
                 "請計算各單元的小計與合計，並請計算配分總和。"):
        para(doc, line, size=11, font=NOTE_FONT, space_after=2)

    # 比例檢查
    if total:
        doc.add_paragraph()
        para(doc, "【認知層次比例檢核】", size=12, bold=True, font=NOTE_FONT, space_after=2)
        for b in range(4):
            pct = round(grand[b][0] / total * 100)
            diff = pct - BLOOM_SUGGEST[b]
            flag = "" if abs(diff) <= 10 else f"　⚠ 偏離建議值 {diff:+d}%"
            para(doc, f"{BLOOM_INTERNAL[b]}：{grand[b][0]} 分（{pct}%）"
                      f"　建議 {BLOOM_SUGGEST[b]}%{flag}",
                 size=11, font=NOTE_FONT, space_after=2)

    path = out_dir / f"{data.get('grade','')}數學第{data.get('exam_number','')}次段考_雙向細目表.docx"
    doc.save(str(path))
    return path


# ══════════════════════════════════════════════════════
# A. 段考卷 — 審題報告
# ══════════════════════════════════════════════════════

def build_review_report(data, out_dir):
    report = data.get("review_report")
    if not report:
        return None
    doc = Document()
    setup_page(doc, 21.0, 29.7, 1.5, lr=2.0)
    para(doc, f"{data.get('grade','')} 數學第{data.get('exam_number','')}次段考　審題報告",
         size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font=NOTE_FONT)
    para(doc, f"範圍：{data.get('scope','')}", size=11,
         align=WD_ALIGN_PARAGRAPH.CENTER, font=NOTE_FONT)
    doc.add_paragraph()

    def table_from(rows, headers, keys, widths=None):
        tb = doc.add_table(rows=len(rows) + 1, cols=len(headers))
        tb.style = "Table Grid"
        for c, h in enumerate(headers):
            set_cell_text(tb.rows[0].cells[c], h, 10, True)
            shade_cell(tb.rows[0].cells[c])
        for r, item in enumerate(rows, start=1):
            for c, k in enumerate(keys):
                val = item.get(k, "")
                set_cell_text(tb.rows[r].cells[c], plain_text(val) if isinstance(val, str) else val,
                              9, align=WD_ALIGN_PARAGRAPH.LEFT if c in (1, len(keys) - 1)
                              else WD_ALIGN_PARAGRAPH.CENTER)
        if widths:
            for c, w in enumerate(widths):
                for row in tb.rows:
                    row.cells[c].width = Cm(w)
        return tb

    if report.get("mc_rows"):
        para(doc, "一、選擇題逐題審查", size=13, bold=True, font=NOTE_FONT)
        table_from(report["mc_rows"],
                   ["題號", "題目摘要", "節次", "判定層次", "答案", "判定理由"],
                   ["number", "summary", "section", "bloom_level", "answer", "reason"],
                   [1.2, 4.2, 2.6, 2.4, 1.2, 5.0])
        doc.add_paragraph()

    if report.get("open_rows"):
        para(doc, "二、非選擇題逐題審查", size=13, bold=True, font=NOTE_FONT)
        rows = [dict(r, big=f"{r.get('big_number','')}{r.get('label','')}")
                for r in report["open_rows"]]
        table_from(rows,
                   ["大題", "題目摘要", "節次", "判定層次", "配分", "判定理由"],
                   ["big", "summary", "section", "bloom_level", "points", "reason"],
                   [1.4, 4.0, 2.6, 2.4, 1.2, 5.0])
        doc.add_paragraph()

    if report.get("section_stats"):
        para(doc, "三、各節次配分統計", size=13, bold=True, font=NOTE_FONT)
        table_from(report["section_stats"],
                   ["節次", "選擇題配分", "非選配分", "合計"],
                   ["section", "mc_pts", "open_pts", "total"])
        doc.add_paragraph()

    if report.get("bloom_stats"):
        para(doc, "四、認知層次分布", size=13, bold=True, font=NOTE_FONT)
        bs = report["bloom_stats"]
        tb = doc.add_table(rows=len(bs) + 1, cols=5)
        tb.style = "Table Grid"
        for c, h in enumerate(["層次", "配分", "佔比", "建議", "差異"]):
            set_cell_text(tb.rows[0].cells[c], h, 10, True)
            shade_cell(tb.rows[0].cells[c])
        for r, (level, v) in enumerate(bs.items(), start=1):
            pct, sug = v.get("pct", 0), v.get("suggest", 0)
            diff = pct - sug
            set_cell_text(tb.rows[r].cells[0], level, 10)
            set_cell_text(tb.rows[r].cells[1], v.get("pts", 0), 10)
            set_cell_text(tb.rows[r].cells[2], f"{pct}%", 10)
            set_cell_text(tb.rows[r].cells[3], f"{sug}%", 10)
            set_cell_text(tb.rows[r].cells[4],
                          f"{diff:+d}%" + ("　⚠" if abs(diff) > 10 else ""), 10)
        doc.add_paragraph()

    if report.get("quality_issues"):
        para(doc, "五、品質問題", size=13, bold=True, font=NOTE_FONT)
        for issue in report["quality_issues"]:
            mark = {"error": "❌", "warning": "⚠️"}.get(issue.get("type"), "・")
            para(doc, f"{mark} {issue.get('title','')}", size=11, bold=True,
                 font=NOTE_FONT, space_after=2)
            para(doc, f"　　{issue.get('description','')}", size=11,
                 font=NOTE_FONT, space_after=6)

    path = out_dir / f"{data.get('grade','')}數學第{data.get('exam_number','')}次段考_審題報告.docx"
    doc.save(str(path))
    return path


# ══════════════════════════════════════════════════════
# B. 情境非選題
# ══════════════════════════════════════════════════════

def build_context_questions(data, out_path, figs):
    doc = Document()
    setup_page(doc, 21.0, 29.7, 1.5, lr=2.0)

    total = sum(q.get("total_points", 0) for q in data.get("questions", []))
    para(doc, data.get("title", "國中數學非選擇題"), size=18, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, font=NOTE_FONT)
    para(doc, f"{data.get('grade','')}　{data.get('unit','')}　滿分 {total} 分",
         size=11, align=WD_ALIGN_PARAGRAPH.CENTER, font=NOTE_FONT)
    para(doc, "班級 _______　座號 _______　姓名 ________________",
         size=11, align=WD_ALIGN_PARAGRAPH.CENTER, font=NOTE_FONT)
    doc.add_paragraph()

    for q in data.get("questions", []):
        n = q["number"]
        hp = para(doc, "", space_after=3)
        set_run_font(hp.add_run(f"第 {n} 題　（{q.get('total_points','')} 分）"),
                     NOTE_FONT, 13, True)
        sp = para(doc, "", space_after=4)
        add_rich_text(sp, q.get("source", ""), size=BODY_SIZE)
        figs.add(doc, q, f"q{n}_main", width_cm=6.0)

        for key, default_label in (("sub1", "(1)"), ("sub2", "(2)")):
            sub = q.get(key)
            if not sub:
                continue
            p = para(doc, "", space_after=3)
            set_run_font(p.add_run(f"{default_label}（{sub.get('points','')} 分）"),
                         BODY_FONT, BODY_SIZE, True)
            add_rich_text(p, sub.get("question", ""), size=BODY_SIZE)
            figs.add(doc, sub, f"q{n}_{key}", width_cm=5.5)
            for _ in range(4):
                para(doc, "", space_after=0)
        doc.add_paragraph()

    # 解析（另起一頁）
    doc.add_page_break()
    para(doc, "解析", size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font=NOTE_FONT)
    doc.add_paragraph()
    for q in data.get("questions", []):
        hp = para(doc, "", space_after=3)
        set_run_font(hp.add_run(f"第 {q['number']} 題"), NOTE_FONT, 13, True)
        for key, default_label in (("sub1", "(1)"), ("sub2", "(2)")):
            sub = q.get(key)
            if not sub:
                continue
            p = para(doc, "", space_after=2)
            set_run_font(p.add_run(f"{default_label} 答："), BODY_FONT, 11, True)
            add_rich_text(p, sub.get("answer", ""), size=11)
            if sub.get("solution"):
                lp = para(doc, "", space_after=2)
                set_run_font(lp.add_run("　　解："), BODY_FONT, 11, True)
                add_rich_text(lp, sub["solution"], size=11)
            tag = " ".join(x for x in (sub.get("bloom_level"), sub.get("bloom_reason")) if x)
            if tag:
                para(doc, f"　　（{tag}）", size=9, font=NOTE_FONT, space_after=6)
        doc.add_paragraph()

    doc.save(str(out_path))
    return out_path


# ══════════════════════════════════════════════════════
# CLI
# ══════════════════════════════════════════════════════

def main():
    ap = argparse.ArgumentParser(description="段考試卷 / 情境非選題 Word 產生器")
    ap.add_argument("json_path", help="題目資料 JSON")
    ap.add_argument("output", help="輸出目錄；情境非選題可直接給 .docx 路徑")
    ap.add_argument("--blueprint-only", action="store_true",
                    help="只產雙向細目表與審題報告（審題模式）")
    ap.add_argument("--figures-dir", default=None,
                    help="幾何圖 PNG 目錄，指定後自動插圖")
    args = ap.parse_args()

    src = Path(args.json_path)
    if not src.exists():
        sys.exit(f"[ERROR] 找不到 JSON：{src}")
    data = json.loads(src.read_text(encoding="utf-8"))
    figs = Figures(args.figures_dir)
    produced = []

    is_context = "questions" in data and not data.get("mc_questions") and not data.get("open_questions")

    if is_context:
        out = Path(args.output)
        if out.suffix.lower() != ".docx":
            out.mkdir(parents=True, exist_ok=True)
            unit = re.sub(r'[\\/:*?"<>|]', "_", str(data.get("unit", "非選擇題")))
            out = out / f"非選擇題_{unit}.docx"
        else:
            out.parent.mkdir(parents=True, exist_ok=True)
        produced.append(build_context_questions(data, out, figs))
    else:
        out_dir = Path(args.output)
        if out_dir.suffix.lower() == ".docx":
            out_dir = out_dir.parent
        out_dir.mkdir(parents=True, exist_ok=True)
        if not args.blueprint_only:
            produced.append(build_exam_paper(data, out_dir, figs))
            produced.append(build_answer_paper(data, out_dir))
        produced.append(build_blueprint(data, out_dir))
        report = build_review_report(data, out_dir)
        if report:
            produced.append(report)

    if figs.missing:
        print(f"[WARN] 有 {len(figs.missing)} 張圖在 --figures-dir 找不到，"
              f"該題未插圖：{', '.join(figs.missing)}")
    for p in produced:
        print(f"[OK] {p}")
    print(f"[DONE] 共產出 {len(produced)} 份文件")


if __name__ == "__main__":
    main()
